from docx import Document
#from docx.shared import Inches
from docx.shared import Cm

from fake_contact import fake_contact

from faker import Faker 
from faker.providers import lorem
fake = Faker('fr_FR')

document = Document()

section = document.sections[-1] 
section.left_margin = Cm(1)
section.right_margin = Cm(1)
section.top_margin = Cm(1)
section.bottom_margin = Cm(1)

document.add_heading('Essai de docx', 0)

p = document.add_paragraph( fake.text() )

p = document.add_paragraph( 'Attributs de texte => ')
p.add_run('on peut ajouter du texte : ')
p.add_run('en GRAS').bold = True
p.add_run(' et du texte : ')
p.add_run('en ITALIQUE.').italic = True

data = [ ]
for x in range(0,10):
    data.append( fake_contact() )

table = document.add_table(rows=1, cols=5)
table.autofit = True
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nom'
hdr_cells[1].text = 'Date de naissance'
hdr_cells[2].text = 'Adresse'
hdr_cells[3].text = 'Téléphone'
hdr_cells[4].text = 'Commentaire'
for n, d, a, t, c in data:
    row_cells = table.add_row().cells
    row_cells[0].text = n
    row_cells[1].text = d.strftime("%d/%m/%Y")
    row_cells[2].text = a
    row_cells[3].text = t
    row_cells[4].text = c

document.add_page_break()

document.save('document1.docx')
